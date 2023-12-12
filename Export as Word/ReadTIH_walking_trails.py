import requests
import json
from docx import Document
from bs4 import BeautifulSoup

api_url = "https://api.stb.gov.sg/content/common/v2/search"
api_key = ""
query_params = {
    "dataset": "walking_trails",
    "limit": 50,
    "offset": 0,
}

headers = {
    "X-API-KEY": api_key
}

# Initialize an empty dictionary to store data grouped by type
grouped_data = {}

while True:
    # Make the API request with the specified query parameters and headers
    response = requests.get(api_url, params=query_params, headers=headers)

    if response.status_code == 200:
        data = response.json()
        if not data['data']:
            break  # No more data to fetch
        else:
            for item in data['data']:
                # Get the "type" of the walking_trail
                walking_trail_type = item.get("type", "Unknown")

                # Create a list for the type if it doesn't exist in the grouped_data dictionary
                if walking_trail_type not in grouped_data:
                    grouped_data[walking_trail_type] = []

                # Remove HTML tags from the description and details
                details = BeautifulSoup(item['body'], "html.parser").get_text()
                item['body'] = details

                # Append the current item to the list for its type
                grouped_data[walking_trail_type].append(item)

            query_params["offset"] += query_params["limit"]  # Update the offset to fetch the next page
    else:
        print(f"Error: {response.status_code} - {response.text}")
        break

# Initialize a Word document
doc = Document()

# Iterate through the grouped data
for walking_trail_type, type_items in grouped_data.items():
    # Add a section header for the walking_trail type
    doc.add_heading(f"Walking Trails Type: {walking_trail_type}", level=1)

    # Iterate through the walking_trails of the current type
    for item in type_items:
        doc.add_heading(item['name'], level=2)
        doc.add_paragraph(f"Type: {item['type']}")
        doc.add_paragraph(f"Description: {item['description']}")
        doc.add_paragraph(f"Details: {item['body']}")
        doc.add_paragraph(f"Official Website: {item['officialWebsite']}")

        doc.add_paragraph(f"Contact: Primary Contact No.: {item.get('contact').get('primaryContactNo')}, Secondary Contact No.: {item.get('contact').get('secondaryContactNo')}")

        doc.add_page_break()

# Save the Word document
doc.save('walking_trails_grouped.docx')

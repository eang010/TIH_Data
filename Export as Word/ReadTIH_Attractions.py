import requests
from docx import Document

api_url = "https://api.stb.gov.sg/content/common/v2/search"
api_key = ""

query_params = {
    "dataset": "attractions",
    "limit": 50,
    "offset": 0,
}

headers = {
    "X-API-KEY": api_key
}

# Initialize a list to store attraction data
attraction_data = []

while True:
    # Make the API request with the specified query parameters and headers
    response = requests.get(api_url, params=query_params, headers=headers)

    if response.status_code == 200:
        data = response.json()
        if not data['data']:
            break  # No more data to fetch
        else:
            for item in data['data']:
                attraction_data.append(item)

            query_params["offset"] += query_params["limit"]  # Update the offset to fetch the next page
    else:
        print(f"Error: {response.status_code} - {response.text}")
        break

# Sort attraction_data by rating in descending order
attraction_data = sorted(attraction_data, key=lambda x: x['rating'], reverse=True)

# Initialize a Word document
doc = Document()

# Iterate through the attractions
for item in attraction_data:
    doc.add_heading(item['name'], level=1)
    doc.add_paragraph(f"Description: {item['description']}")
    doc.add_paragraph(f"Rating: {item['rating']}")
    doc.add_paragraph(f"Nearest MRT Station: {item['nearestMrtStation']}")
    doc.add_paragraph(f"Official Website: {item['officialWebsite']}")
    doc.add_paragraph(f"Admission Info: {item['admissionInfo']}")
    doc.add_paragraph(f"Pricing: {item.get('pricing', {}).get('others')}")
    doc.add_paragraph(f"Amenities: {item['amenities']}")

    # Combine the address fields into a single sentence
    address_parts = [item['address']['block'], item['address']['streetName'], item['address']['postalCode']]
    address_sentence = ' '.join(part for part in address_parts if part)
    doc.add_paragraph(f"Address: {address_sentence}")

    doc.add_paragraph("Business Hours:")
    for business_hour in item.get("businessHour", []):
        doc.add_paragraph(f"{business_hour['day']} - Open: {business_hour['openTime']}, Close: {business_hour['closeTime']}, Description: {business_hour['description']}")

    doc.add_paragraph(f"Business Hour Notes: {item.get('businessHourNotes', {}).get('notes')}")

    doc.add_page_break()

# Save the Word document
doc.save('attractions_sorted_by_rating.docx')

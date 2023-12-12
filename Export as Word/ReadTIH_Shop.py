import requests
import json
from docx import Document

api_url = "https://api.stb.gov.sg/content/common/v2/search"
api_key = ""
query_params = {
    "dataset": "shops",
    "limit": 50,
    "offset": 0,
}

headers = {
    "X-API-KEY": api_key
}

# Initialize an empty dictionary to store data grouped by type
grouped_data = {}

while True:
    # Make the API request with the specified query parameters and headers
    response = requests.get(api_url, params=query_params, headers=headers)

    if response.status_code == 200:
        data = response.json()
        if not data['data']:
            break  # No more data to fetch
        else:
            for item in data['data']:
                # Get the "type" of the shops
                shops_type = item.get("type", "Unknown")

                # Create a list for the type if it doesn't exist in the grouped_data dictionary
                if shops_type not in grouped_data:
                    grouped_data[shops_type] = []

                # Append the current item to the list for its type
                grouped_data[shops_type].append(item)

            query_params["offset"] += query_params["limit"]  # Update the offset to fetch the next page
    else:
        print(f"Error: {response.status_code} - {response.text}")
        break

# Initialize a Word document
doc = Document()

# Iterate through the grouped data
for shops_type, type_items in grouped_data.items():
    # Add a section header for the shops type
    doc.add_heading(f"Shops Type: {shops_type}", level=1)

    # Iterate through the shops of the current type
    for item in type_items:
        doc.add_heading(item['name'], level=2)
        doc.add_paragraph(f"Description: {item['description']}")
        doc.add_paragraph(f"Nearest MRT Station: {item['nearestMrtStation']}")
        doc.add_paragraph(f"Official Website: {item['officialWebsite']}")
        doc.add_paragraph(f"Amenities: {item['amenities']}")

        # Combine the address fields into a single sentence
        address_parts = [item['address']['block'], item['address']['streetName'], item['address']['postalCode']]
        address_sentence = ' '.join(part for part in address_parts if part)
        doc.add_paragraph(f"Address: {address_sentence}")

        doc.add_paragraph("Business Hours:")
        for business_hour in item.get("businessHour", []):
            doc.add_paragraph(f"{business_hour['day']} - Open: {business_hour['openTime']}, Close: {business_hour['closeTime']}, Description: {business_hour['description']}")

        doc.add_paragraph(f"Business Hour Notes: {item.get('businessHourNotes', {}).get('notes')}")

        doc.add_paragraph(f"Contact: Primary Contact No.: {item.get('contact').get('primaryContactNo')}, Secondary Contact No.: {item.get('contact').get('secondaryContactNo')}")
        doc.add_paragraph(f"Location: Latitude - {item.get('location').get('latitude')}, Longitude - {item.get('location').get('longitude')}")

        doc.add_page_break()

# Save the Word document
doc.save('shops_grouped.docx')
